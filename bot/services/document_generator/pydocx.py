import asyncio
import io
import logging
import os
from pathlib import Path

from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.shared import Pt
from docxtpl import DocxTemplate

logger = logging.getLogger(__name__)

OUTPUT_DIR = Path(__file__).resolve().parents[3] / "data" / "history_of_illness" / "generated"
PROJECT_ROOT = Path(__file__).resolve().parents[3]

TEETH_TABLE_INDEX = 1

TOOTH_MAP = {
    # Верхняя челюсть
    18: (0, 0), 17: (0, 1), 16: (0, 2), 15: (0, 3),
    14: (0, 4), 13: (0, 5), 12: (0, 6), 11: (0, 7),
    21: (0, 8), 22: (0, 9), 23: (0, 10), 24: (0, 11),
    25: (0, 12), 26: (0, 13), 27: (0, 14), 28: (0, 15),

    # Нижняя челюсть
    48: (3, 0), 47: (3, 1), 46: (3, 2), 45: (3, 3),
    44: (3, 4), 43: (3, 5), 42: (3, 6), 41: (3, 7),
    31: (3, 8), 32: (3, 9), 33: (3, 10), 34: (3, 11),
    35: (3, 12), 36: (3, 13), 37: (3, 14), 38: (3, 15),
}


def mark_tooth(table, tooth: int, marker: str) -> None:
    if tooth not in TOOTH_MAP:
        return

    row, col = TOOTH_MAP[tooth]

    cell = table.cell(row, col)
    cell.text = ""
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER

    paragraph = cell.paragraphs[0]
    paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    run = paragraph.add_run(marker)
    run.bold = True
    run.font.size = Pt(12)


def mark_teeth(table, tooth_map: list[dict]) -> None:
    if not isinstance(tooth_map, list):
        logger.warning("tooth_map не является списком, пропускаю отметку зубов: %r", tooth_map)
        return

    for entry in tooth_map:
        try:
            tooth = entry["tooth"]
            marker = entry["marker"]
        except (TypeError, KeyError):
            logger.warning("Некорректная запись в tooth_map, пропускаю: %r", entry)
            continue

        if not isinstance(tooth, int) or not isinstance(marker, str):
            logger.warning("Некорректная запись в tooth_map, пропускаю: %r", entry)
            continue

        mark_tooth(table, tooth, marker)


def _render_docx_sync(data: dict, tooth_map: list[dict], output_path: Path, template_path: Path) -> None:
    template = DocxTemplate(str(template_path))
    template.render(data)

    buffer = io.BytesIO()
    template.save(buffer)
    buffer.seek(0)

    doc = Document(buffer)

    try:
        teeth_table = doc.tables[TEETH_TABLE_INDEX]
    except IndexError:
        logger.warning("Шаблон не содержит таблицу зубной карты (индекс %s), пропускаю отметку зуба.", TEETH_TABLE_INDEX)
    else:
        mark_teeth(teeth_table, tooth_map)

    logger.info("Отмечены зубы: %s", tooth_map)

    output_path.parent.mkdir(parents=True, exist_ok=True)
    tmp_path = output_path.with_name(output_path.name + ".tmp")

    try:
        doc.save(str(tmp_path))
        os.replace(tmp_path, output_path)
    except Exception:
        if tmp_path.exists():
            tmp_path.unlink(missing_ok=True)
        raise

    logger.info("Документ сохранён: %s", output_path)


async def create_docx(data: dict, tooth_map: list[dict], output_path: Path, template_path: str) -> str:
    """Render the medical record template and mark the affected teeth.

    `data` must already contain every template placeholder (appointment_date,
    full_name, gender, birth_date, phone, diagnosis, complaints, diseases,
    examination, treatment) with LLM-key-to-template-key mapping already
    applied by the caller. `tooth_map` is a list of dicts shaped
    `{"tooth": int, "marker": str}`, sourced from the LLM's structured
    output; it is an empty list when the diagnosis mentions no teeth.
    `output_path` is the absolute destination path built by the caller (see
    MedicalRecordService._build_output_path). `template_path` is the
    per-clinic-instance .docx template path resolved by the caller (see
    MEDICAL_RECORD_TEMPLATE_BY_INSTANCE in bot.config.clinic_instances).

    The render is saved to a temp file next to output_path and moved into
    place with os.replace, which is atomic on both POSIX and Windows, so a
    caller reading output_path never observes a half-written document.

    Returns the path to the generated .docx file.
    """
    resolved_template_path = PROJECT_ROOT / template_path

    await asyncio.to_thread(_render_docx_sync, data, tooth_map, output_path, resolved_template_path)

    return str(output_path)
